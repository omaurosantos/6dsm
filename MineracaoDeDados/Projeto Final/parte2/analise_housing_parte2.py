from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


SCRIPT_DIR = Path(__file__).resolve().parent
DATASET = SCRIPT_DIR / "Housing.csv"
REPORT = SCRIPT_DIR / "Relatorio_Projeto_Final_Parte2_Housing.docx"
GRAPH = SCRIPT_DIR / "kmeans_grupos.png"
RANDOM_SEED = 42


YES_NO_COLUMNS = [
    "mainroad",
    "guestroom",
    "basement",
    "hotwaterheating",
    "airconditioning",
    "prefarea",
]
FURNISHING_MAP = {"unfurnished": 0, "semi-furnished": 1, "furnished": 2}
CLASS_LABELS = ["BaixoPadrão", "MédioPadrão", "AltoPadrão"]


def format_number(value: float, decimals: int = 2) -> str:
    return f"{value:,.{decimals}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def load_and_preprocess() -> tuple[pd.DataFrame, dict[str, int], pd.DataFrame]:
    original = pd.read_csv(DATASET)
    missing = original.isna().sum()
    missing_summary = {column: int(count) for column, count in missing.items() if count > 0}

    df = original.copy()
    for column in df.columns:
        if df[column].isna().sum() == 0:
            continue
        if pd.api.types.is_numeric_dtype(df[column]):
            df[column] = df[column].fillna(df[column].mean())
        else:
            df[column] = df[column].fillna(df[column].mode()[0])

    for column in YES_NO_COLUMNS:
        df[column] = df[column].map({"yes": 1, "no": 0})
    df["furnishingstatus"] = df["furnishingstatus"].map(FURNISHING_MAP)

    menor = df["price"].min()
    maior = df["price"].max()
    intervalo = (maior - menor) / 3
    valor1 = menor + intervalo
    valor2 = menor + (2 * intervalo)
    df["ClassePadrao"] = pd.cut(
        df["price"],
        bins=[-np.inf, valor1, valor2, np.inf],
        labels=CLASS_LABELS,
        include_lowest=True,
    )
    return df, missing_summary, original


def train_test_split(n: int, test_size: float = 0.2) -> tuple[np.ndarray, np.ndarray]:
    rng = np.random.default_rng(RANDOM_SEED)
    indices = np.arange(n)
    rng.shuffle(indices)
    test_n = int(round(n * test_size))
    return indices[test_n:], indices[:test_n]


def linear_regression(df: pd.DataFrame) -> dict[str, object]:
    features = [column for column in df.columns if column not in ["price", "ClassePadrao"]]
    x = df[features].to_numpy(dtype=float)
    y = df["price"].to_numpy(dtype=float)
    train_idx, test_idx = train_test_split(len(df), 0.2)

    x_train = np.column_stack([np.ones(len(train_idx)), x[train_idx]])
    x_test = np.column_stack([np.ones(len(test_idx)), x[test_idx]])
    y_train = y[train_idx]
    y_test = y[test_idx]

    coef = np.linalg.pinv(x_train.T @ x_train) @ x_train.T @ y_train
    predictions = x_test @ coef

    residuals = y_test - predictions
    ss_res = float(np.sum(residuals**2))
    ss_tot = float(np.sum((y_test - y_test.mean()) ** 2))
    r2 = 1 - ss_res / ss_tot
    mae = float(np.mean(np.abs(residuals)))
    rmse = float(np.sqrt(np.mean(residuals**2)))

    feature_std = df[features].std(ddof=0).to_numpy(dtype=float)
    y_std = df["price"].std(ddof=0)
    standardized = coef[1:] * feature_std / y_std
    influence = sorted(
        zip(features, coef[1:], standardized),
        key=lambda item: abs(item[2]),
        reverse=True,
    )

    return {
        "features": features,
        "intercept": coef[0],
        "coefficients": dict(zip(features, coef[1:])),
        "r2": r2,
        "mae": mae,
        "rmse": rmse,
        "influence": influence,
    }


def standardize(x: np.ndarray) -> tuple[np.ndarray, np.ndarray, np.ndarray]:
    mean = x.mean(axis=0)
    std = x.std(axis=0)
    std[std == 0] = 1
    return (x - mean) / std, mean, std


def stratified_folds(y: np.ndarray, folds: int = 10) -> list[np.ndarray]:
    rng = np.random.default_rng(RANDOM_SEED)
    fold_indices = [[] for _ in range(folds)]
    for label in np.unique(y):
        indices = np.where(y == label)[0]
        rng.shuffle(indices)
        for index, value in enumerate(indices):
            fold_indices[index % folds].append(int(value))
    return [np.array(sorted(fold), dtype=int) for fold in fold_indices]


def confusion_matrix(y_true: np.ndarray, y_pred: np.ndarray, labels: list[str]) -> np.ndarray:
    lookup = {label: index for index, label in enumerate(labels)}
    matrix = np.zeros((len(labels), len(labels)), dtype=int)
    for true, pred in zip(y_true, y_pred):
        matrix[lookup[true], lookup[pred]] += 1
    return matrix


class LDAClassifier:
    def fit(self, x: np.ndarray, y: np.ndarray) -> "LDAClassifier":
        self.labels = np.unique(y)
        self.means = {}
        self.priors = {}
        pooled = np.zeros((x.shape[1], x.shape[1]))
        for label in self.labels:
            group = x[y == label]
            self.means[label] = group.mean(axis=0)
            self.priors[label] = len(group) / len(x)
            centered = group - self.means[label]
            pooled += centered.T @ centered
        pooled /= max(len(x) - len(self.labels), 1)
        self.inv_cov = np.linalg.pinv(pooled + np.eye(x.shape[1]) * 1e-6)
        return self

    def predict(self, x: np.ndarray) -> np.ndarray:
        scores = []
        for label in self.labels:
            mean = self.means[label]
            score = x @ self.inv_cov @ mean - 0.5 * mean.T @ self.inv_cov @ mean + np.log(self.priors[label])
            scores.append(score)
        return self.labels[np.argmax(np.column_stack(scores), axis=1)]


@dataclass
class TreeNode:
    prediction: str
    feature: int | None = None
    threshold: float | None = None
    left: "TreeNode | None" = None
    right: "TreeNode | None" = None


class DecisionTreeClassifier:
    def __init__(self, max_depth: int = 5, min_samples_split: int = 12):
        self.max_depth = max_depth
        self.min_samples_split = min_samples_split

    def fit(self, x: np.ndarray, y: np.ndarray) -> "DecisionTreeClassifier":
        self.labels = np.unique(y)
        self.root = self._build(x, y, 0)
        return self

    def _gini(self, y: np.ndarray) -> float:
        _, counts = np.unique(y, return_counts=True)
        probabilities = counts / len(y)
        return float(1 - np.sum(probabilities**2))

    def _majority(self, y: np.ndarray) -> str:
        labels, counts = np.unique(y, return_counts=True)
        return str(labels[np.argmax(counts)])

    def _best_split(self, x: np.ndarray, y: np.ndarray) -> tuple[int | None, float | None]:
        best_feature, best_threshold, best_score = None, None, float("inf")
        for feature in range(x.shape[1]):
            values = np.unique(x[:, feature])
            if len(values) <= 1:
                continue
            thresholds = (values[:-1] + values[1:]) / 2
            if len(thresholds) > 40:
                thresholds = np.quantile(values, np.linspace(0.05, 0.95, 40))
            for threshold in thresholds:
                left = y[x[:, feature] <= threshold]
                right = y[x[:, feature] > threshold]
                if len(left) == 0 or len(right) == 0:
                    continue
                score = (len(left) * self._gini(left) + len(right) * self._gini(right)) / len(y)
                if score < best_score:
                    best_feature, best_threshold, best_score = feature, float(threshold), score
        return best_feature, best_threshold

    def _build(self, x: np.ndarray, y: np.ndarray, depth: int) -> TreeNode:
        node = TreeNode(prediction=self._majority(y))
        if depth >= self.max_depth or len(y) < self.min_samples_split or len(np.unique(y)) == 1:
            return node
        feature, threshold = self._best_split(x, y)
        if feature is None:
            return node
        mask = x[:, feature] <= threshold
        node.feature = feature
        node.threshold = threshold
        node.left = self._build(x[mask], y[mask], depth + 1)
        node.right = self._build(x[~mask], y[~mask], depth + 1)
        return node

    def _predict_one(self, row: np.ndarray) -> str:
        node = self.root
        while node.feature is not None:
            node = node.left if row[node.feature] <= node.threshold else node.right
        return node.prediction

    def predict(self, x: np.ndarray) -> np.ndarray:
        return np.array([self._predict_one(row) for row in x])


class KNNClassifier:
    def __init__(self, k: int = 5):
        self.k = k

    def fit(self, x: np.ndarray, y: np.ndarray) -> "KNNClassifier":
        self.x = x
        self.y = y
        return self

    def predict(self, x: np.ndarray) -> np.ndarray:
        predictions = []
        for row in x:
            distances = np.linalg.norm(self.x - row, axis=1)
            nearest = self.y[np.argsort(distances)[: self.k]]
            labels, counts = np.unique(nearest, return_counts=True)
            predictions.append(labels[np.argmax(counts)])
        return np.array(predictions)


class GaussianNBClassifier:
    def fit(self, x: np.ndarray, y: np.ndarray) -> "GaussianNBClassifier":
        self.labels = np.unique(y)
        self.mean = {}
        self.var = {}
        self.prior = {}
        for label in self.labels:
            group = x[y == label]
            self.mean[label] = group.mean(axis=0)
            self.var[label] = group.var(axis=0) + 1e-6
            self.prior[label] = len(group) / len(x)
        return self

    def predict(self, x: np.ndarray) -> np.ndarray:
        scores = []
        for label in self.labels:
            mean = self.mean[label]
            var = self.var[label]
            log_probability = -0.5 * np.sum(np.log(2 * np.pi * var), axis=0)
            log_probability += -0.5 * np.sum(((x - mean) ** 2) / var, axis=1)
            scores.append(np.log(self.prior[label]) + log_probability)
        return self.labels[np.argmax(np.column_stack(scores), axis=1)]


def run_classification(df: pd.DataFrame) -> dict[str, dict[str, object]]:
    features = [column for column in df.columns if column not in ["price", "ClassePadrao"]]
    x = df[features].to_numpy(dtype=float)
    y = df["ClassePadrao"].astype(str).to_numpy()
    x_scaled, _, _ = standardize(x)
    folds = stratified_folds(y, 10)

    models = {
        "LDA": LDAClassifier,
        "Árvore de Decisão": DecisionTreeClassifier,
        "KNN": KNNClassifier,
        "Naive Bayes": GaussianNBClassifier,
    }
    results: dict[str, dict[str, object]] = {}

    for name, factory in models.items():
        accuracies = []
        matrix = np.zeros((len(CLASS_LABELS), len(CLASS_LABELS)), dtype=int)
        for test_idx in folds:
            train_idx = np.setdiff1d(np.arange(len(y)), test_idx)
            model = factory().fit(x_scaled[train_idx], y[train_idx])
            pred = model.predict(x_scaled[test_idx])
            accuracies.append(float(np.mean(pred == y[test_idx])))
            matrix += confusion_matrix(y[test_idx], pred, CLASS_LABELS)
        results[name] = {
            "accuracy": float(np.mean(accuracies)),
            "folds": accuracies,
            "matrix": matrix,
        }
    return results


def kmeans(x: np.ndarray, k: int = 3, max_iter: int = 200) -> tuple[np.ndarray, np.ndarray]:
    rng = np.random.default_rng(RANDOM_SEED)
    centroids = x[rng.choice(len(x), size=k, replace=False)]
    labels = np.zeros(len(x), dtype=int)
    for _ in range(max_iter):
        distances = np.linalg.norm(x[:, None, :] - centroids[None, :, :], axis=2)
        new_labels = np.argmin(distances, axis=1)
        if np.array_equal(labels, new_labels):
            break
        labels = new_labels
        for group in range(k):
            if np.any(labels == group):
                centroids[group] = x[labels == group].mean(axis=0)
    return labels, centroids


def pca_2d(x: np.ndarray) -> np.ndarray:
    centered = x - x.mean(axis=0)
    _, _, vt = np.linalg.svd(centered, full_matrices=False)
    return centered @ vt[:2].T


def draw_kmeans_graph(points: np.ndarray, labels: np.ndarray, path: Path) -> None:
    width, height = 980, 680
    margin = 70
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()
    colors = [(44, 123, 182), (215, 84, 66), (76, 153, 102)]
    x_values, y_values = points[:, 0], points[:, 1]
    x_min, x_max = x_values.min(), x_values.max()
    y_min, y_max = y_values.min(), y_values.max()

    def scale_x(value: float) -> float:
        return margin + (value - x_min) / (x_max - x_min) * (width - 2 * margin)

    def scale_y(value: float) -> float:
        return height - margin - (value - y_min) / (y_max - y_min) * (height - 2 * margin)

    draw.rectangle([margin, margin, width - margin, height - margin], outline=(170, 170, 170), width=2)
    draw.text((margin, 25), "K-Means com 3 grupos - projecao PCA 2D", fill=(30, 30, 30), font=font)
    for group in range(3):
        group_points = points[labels == group]
        for x, y in group_points:
            cx, cy = scale_x(float(x)), scale_y(float(y))
            draw.ellipse([cx - 4, cy - 4, cx + 4, cy + 4], fill=colors[group], outline=colors[group])
        draw.rectangle([width - 230, 70 + group * 32, width - 210, 90 + group * 32], fill=colors[group])
        draw.text(
            (width - 200, 70 + group * 32),
            f"Grupo {group + 1}: {len(group_points)} casas",
            fill=(30, 30, 30),
            font=font,
        )
    image.save(path)


def run_kmeans(df: pd.DataFrame) -> dict[str, object]:
    numeric = df.drop(columns=["ClassePadrao"])
    x = numeric.to_numpy(dtype=float)
    x_scaled, _, _ = standardize(x)
    labels, centroids = kmeans(x_scaled, 3)
    points = pca_2d(x_scaled)
    draw_kmeans_graph(points, labels, GRAPH)

    summary = []
    for group in range(3):
        group_df = df[labels == group]
        summary.append(
            {
                "Grupo": group + 1,
                "Quantidade": len(group_df),
                "Preço médio": float(group_df["price"].mean()),
                "Área média": float(group_df["area"].mean()),
                "Quartos médios": float(group_df["bedrooms"].mean()),
                "Banheiros médios": float(group_df["bathrooms"].mean()),
            }
        )
    return {"labels": labels, "centroids": centroids, "summary": summary, "graph": GRAPH}


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(31, 77, 120)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)


def build_report(
    df: pd.DataFrame,
    missing_summary: dict[str, int],
    regression: dict[str, object],
    classification: dict[str, dict[str, object]],
    clusters: dict[str, object],
) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Projeto Final - Mineração de Dados")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 77, 120)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Parte 2: Pré-processamento, Regressão, Classificação e Agrupamento").bold = True
    doc.add_paragraph("Base de dados utilizada: Housing.csv")
    doc.add_paragraph(f"Total de registros: {len(df)}")

    add_heading(doc, "4) Pré-processamento dos Dados", 1)
    if missing_summary:
        rows = [[column, str(count)] for column, count in missing_summary.items()]
        doc.add_paragraph("Foram encontrados dados ausentes nos seguintes atributos:")
        add_table(doc, ["Atributo", "Quantidade de ausentes"], rows)
    else:
        doc.add_paragraph("Não foram encontrados dados ausentes na base.")
    doc.add_paragraph(
        "Os atributos categóricos com valores yes/no foram transformados em 1/0. "
        "O atributo furnishingstatus foi transformado em: furnished = 2, semi-furnished = 1, unfurnished = 0."
    )

    add_heading(doc, "5) Regressão Linear", 1)
    add_table(
        doc,
        ["Métrica", "Resultado"],
        [
            ["R²", format_number(float(regression["r2"]), 4)],
            ["MAE", format_number(float(regression["mae"]), 2)],
            ["RMSE", format_number(float(regression["rmse"]), 2)],
        ],
    )
    add_heading(doc, "5.b) Função linear gerada", 2)
    equation_parts = [f"{format_number(float(regression['intercept']), 2)}"]
    for feature, coefficient in regression["coefficients"].items():
        sign = "+" if coefficient >= 0 else "-"
        equation_parts.append(f" {sign} {format_number(abs(float(coefficient)), 2)}*{feature}")
    doc.add_paragraph("price = " + "".join(equation_parts))

    add_heading(doc, "5.c) Variáveis que mais influenciaram o preço", 2)
    rows = [
        [feature, format_number(float(coef), 2), format_number(float(std_coef), 4)]
        for feature, coef, std_coef in regression["influence"][:5]
    ]
    add_table(doc, ["Variável", "Coeficiente", "Coeficiente padronizado"], rows)
    doc.add_paragraph(
        "A influência foi ordenada pelo valor absoluto do coeficiente padronizado, pois isso permite comparar "
        "variáveis em escalas diferentes."
    )

    add_heading(doc, "6) Classificação com Cross Validation = 10 folds", 1)
    rows = [[name, format_number(values["accuracy"] * 100, 2) + "%"] for name, values in classification.items()]
    add_table(doc, ["Algoritmo", "Acurácia média"], rows)
    best_name = max(classification, key=lambda name: classification[name]["accuracy"])
    doc.add_paragraph(
        f"O melhor resultado foi obtido pelo algoritmo {best_name}, com acurácia média de "
        f"{format_number(classification[best_name]['accuracy'] * 100, 2)}%."
    )
    for name, values in classification.items():
        add_heading(doc, f"Matriz de confusão - {name}", 2)
        matrix = values["matrix"]
        rows = []
        for i, true_label in enumerate(CLASS_LABELS):
            rows.append([true_label] + [str(int(matrix[i, j])) for j in range(len(CLASS_LABELS))])
        add_table(doc, ["Real \\ Previsto"] + CLASS_LABELS, rows)

    add_heading(doc, "7) Agrupamento com K-Means", 1)
    rows = [
        [
            str(item["Grupo"]),
            str(item["Quantidade"]),
            format_number(item["Preço médio"], 2),
            format_number(item["Área média"], 2),
            format_number(item["Quartos médios"], 2),
            format_number(item["Banheiros médios"], 2),
        ]
        for item in clusters["summary"]
    ]
    add_table(doc, ["Grupo", "Qtd.", "Preço médio", "Área média", "Quartos", "Banheiros"], rows)
    doc.add_paragraph("Representação gráfica dos 3 grupos gerados pelo K-Means:")
    doc.add_picture(str(clusters["graph"]), width=Inches(6.3))

    doc.save(REPORT)


def main() -> None:
    df, missing_summary, _ = load_and_preprocess()
    regression = linear_regression(df)
    classification = run_classification(df)
    clusters = run_kmeans(df)
    build_report(df, missing_summary, regression, classification, clusters)

    print("Relatório:", REPORT)
    print("Gráfico:", GRAPH)
    print("R2:", regression["r2"])
    print("MAE:", regression["mae"])
    print("RMSE:", regression["rmse"])
    for name, values in classification.items():
        print(name, values["accuracy"])


if __name__ == "__main__":
    main()
